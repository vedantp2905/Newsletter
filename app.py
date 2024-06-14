import os
import asyncio
import requests
from typing import Type, Any
from io import BytesIO
import streamlit as st
from pydantic import BaseModel, Field
from docx import Document
from crewai_tools import ScrapeWebsiteTool
from crewai_tools.tools.base_tool import BaseTool
from langchain_openai import ChatOpenAI
from langchain_google_genai import ChatGoogleGenerativeAI
from crewai import Agent, Task, Crew

# Define SerpApi key (you should set this in your environment or through Streamlit input)
serp_api_key = ''

class SerpApiGoogleSearchToolSchema(BaseModel):
    q: str = Field(..., description="Parameter defines the query you want to search. You can use anything that you would use in a regular Google search. e.g. inurl:, site:, intitle:.")
    tbs: str = Field("qdr:w2", description="Time filter to limit the search to the last two weeks.")

class SerpApiGoogleSearchTool(BaseTool):
    name: str = "Google Search"
    description: str = "Search the internet"
    args_schema: Type[BaseModel] = SerpApiGoogleSearchToolSchema
    search_url: str = "https://serpapi.com/search"
    
    def _run(self, q: str, tbs: str = "qdr:w2", **kwargs: Any) -> Any:
        payload = {
            "engine": "google",
            "q": q,
            "tbs": tbs,
            "api_key": serp_api_key,
        }
        headers = {'content-type': 'application/json'}
        response = requests.get(self.search_url, headers=headers, params=payload)
        results = response.json()
    
        summary = ""
        if 'answer_box_list' in results:
            summary += str(results['answer_box_list'])
        elif 'answer_box' in results:
            summary += str(results['answer_box'])
        elif 'organic_results' in results:
            summary += str(results['organic_results'])
        elif 'sports_results' in results:
            summary += str(results['sports_results'])
        elif 'knowledge_graph' in results:
            summary += str(results['knowledge_graph'])
        elif 'top_stories' in results:
            summary += str(results['top_stories'])
        
        print(summary)
        
        return summary

class SerpApiGoogleImagesSearchToolSchema(BaseModel):
    q: str = Field(..., description="Query for image search. e.g., 'cats', 'sunset'.")
    tbs: str = Field("qdr:w2", description="Time filter to limit the search to the last two weeks.")

class SerpApiGoogleImagesSearchTool(BaseTool):
    name: str = "Google Images Search"
    description: str = "Search the internet for images"
    args_schema: Type[BaseModel] = SerpApiGoogleImagesSearchToolSchema
    search_url: str = "https://serpapi.com/search"
    
    def _run(self, q: str, tbs: str = "qdr:w2", **kwargs: Any) -> Any:
        payload = {
            "engine": "google_images",
            "q": q,
            "tbs": tbs,
            "api_key": serp_api_key,
        }
        headers = {'content-type': 'application/json'}
        response = requests.get(self.search_url, headers=headers, params=payload)
        results = response.json()
    
        summary = ""
        image_urls = []
        if 'images_results' in results:
            for image in results['images_results']:
                summary += f"Title: {image.get('title')}\n"
                summary += f"Source: {image.get('source')}\n"
                summary += f"Link: {image.get('link')}\n"
                summary += f"Thumbnail: {image.get('thumbnail')}\n\n"
                image_urls.append(image.get('thumbnail'))
        
        print(summary)
        
        return summary, image_urls

def generate_text(llm, topic):
    inputs = {'topic': topic}
    
    search_tool = SerpApiGoogleSearchTool()
    search_image = SerpApiGoogleImagesSearchTool()
    
    scrape_tool = ScrapeWebsiteTool(
        name="website_scraper",
        description="""Scrape content from web pages. Action Input should look like this:
                       {"website_url": "<URL of the webpage to scrape>"}""",
    )
    
    researcher_agent = Agent(
        role='Newsletter Content Researcher',
        goal='Search the latest top 5 stories on the given topic, find unique 5 URLs containing the stories, relevant image for each story, and scrape relevant information from these URLs.',
        backstory=(
            "An experienced researcher with strong skills in web scraping, fact-finding, and "
            "analyzing recent trends to provide up-to-date information for high-quality newsletters."
        ),
        verbose=True,
        allow_delegation=False,
        llm=llm
    )

    writer_agent = Agent(
        role='Content Writer',
        goal='Write detailed, engaging, and informative summaries of the developments found by the researcher using the format specified.',
        backstory=("An experienced writer with a background in journalism and content creation. "
                   "Skilled in crafting compelling narratives and distilling complex information into "
                   "accessible formats. Adept at conducting research and synthesizing insights for engaging content."),
        verbose=True,
        allow_delegation=False,
        llm=llm
    )

    reviewer_agent = Agent(
        role='Content Reviewer',
        goal='Review and refine content drafts to ensure they meet high standards of quality and impact like major newsletters.',
        backstory=("A meticulous reviewer with extensive experience in editing and proofreading, "
                   "known for their keen eye for detail and commitment to maintaining the highest quality standards in published content."),
        verbose=True,
        allow_delegation=False,
        llm=llm
    )
    
    final_writer_agent = Agent(
        role='Final Content Writer',
        goal='Compile, refine, and structure all reviewed and approved content into a cohesive and engaging newsletter format. Ensure that the final product is polished, logically structured, and ready for publication, providing a seamless and informative reading experience for the audience.',
        backstory=("An accomplished writer and editor with extensive experience in journalism, content creation, and editorial management. "
                   "Known for their ability to craft compelling narratives and ensure consistency and quality across all sections of a publication. "
                   "With a keen eye for detail and a deep understanding of audience engagement, this writer excels in transforming raw content into polished, professional-grade newsletters that captivate readers and deliver clear, valuable insights."),
        verbose=True,
        allow_delegation=False,
        llm=llm
    )
    
    task_researcher = Task(
        description=(f'Research and identify the most interesting 5 stories on the topic of {topic} '
                     'Scrape detailed content from relevant websites to gather comprehensive material.'
                     'Find relevant image for each story'),
        agent=researcher_agent,
        expected_output=('A list of 3-4 recent developments and 2 stories from more than a week ago with their respective website URLs. '
                         'Scraped content from all URLs that can be used further by the writer.'
                         'A relevant image for each story'),
        tools=[search_tool, scrape_tool, search_image]
    )

    task_writer = Task(
        description=('Write detailed summaries of the recent developments identified by the researcher. '
                     'Ensure each summary is informative, engaging, and provides clear insights into the development.'),
        agent=writer_agent,
        expected_output=('Summarized content for all the stories, each summary being 150-200 words long, '
                         'with clear and concise information.')
    )

    task_reviewer = Task(
        description=('Review the summarized content provided by the writer for accuracy, coherence, and quality. '
                     'Ensure that the content is free from errors and meets the publication standards.'),
        agent=reviewer_agent,
        expected_output=('Reviewed content with suggestions for improvements, if any. '
                         'Final versions of summaries that are ready for inclusion in the newsletter.')
    )

    task_final_writer = Task(
        description=('Compile the reviewed and refined content into a well-structured newsletter format. '
                     'Ensure the newsletter is visually appealing and flows logically from one section to the next.'),
        agent=final_writer_agent,
        expected_output=(
            """Final newsletter document with all the reviewed summaries, formatted and ready for publication. 
            The newsletter should include:
            - Introduction:
                - A compelling hook sentence to engage readers and encourage them to read the entire newsletter.
            - Contents section:
                - Summarize each story or development in one interesting sentence.
            - Main content sections (5-6 developments/stories):
                - Each story should have:
                    - Image after title
                    - A small introduction.
                    - Main details presented in 3-4 bullet points.
                    - Explanation of why it matters or a call to action
            - Conclusion:
                - Wrap up the newsletter by summarizing all content and providing a final thought or conclusion.
            """
        )
    )

    crew = Crew(
        agents=[researcher_agent, writer_agent, reviewer_agent, final_writer_agent],
        tasks=[task_researcher, task_writer, task_reviewer, task_final_writer],
        verbose=2,
        context={"Blog Topic is ": topic}
    )

    result = crew.kickoff(inputs=inputs)

    return result

# Streamlit web application
def main():
    st.header('AI Newsletter Content Generator')
    global serp_api_key
    
    with st.sidebar:
        with st.form('Gemini/OpenAI'):
            model = st.radio('Choose Your LLM', ('Gemini', 'OpenAI'))
            api_key = st.text_input(f'Enter your API key', type="password")
            serp_api_key = st.text_input(f'Enter your SerpAPI key', type="password")
            submitted = st.form_submit_button("Submit")

    if api_key and serp_api_key:
        if model == 'OpenAI':
            async def setup_openai():
                os.environ["OPENAI_API_KEY"] = api_key
                llm = ChatOpenAI(temperature=0.6, max_tokens=3500)
                print("Configured OpenAI model")
                return llm

            llm = asyncio.run(setup_openai())

        elif model == 'Gemini':
            async def setup_gemini():
                llm = ChatGoogleGenerativeAI(
                    model="gemini-1.5-flash",
                    verbose=True,
                    temperature=0.6,
                    google_api_key=api_key
                )
                print("Gemini Configured")
                return llm

            llm = asyncio.run(setup_gemini())
        
        topic = st.text_input("Enter the blog topic:")

        if st.button("Generate Newsletter Content"):
            with st.spinner("Generating content..."):
                generated_content, image_urls = generate_text(llm, topic)
                
                content_lines = generated_content.split('\n')
                first_line = content_lines[0]
                remaining_content = '\n'.join(content_lines[1:])

                st.markdown(first_line)
                st.markdown(remaining_content)

                # Display images
                for image_url in image_urls:
                    st.image(image_url)

                doc = Document()
                doc.add_heading(topic, 0)
                doc.add_paragraph(first_line)
                doc.add_paragraph(remaining_content)

                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                st.download_button(
                    label="Download as Word Document",
                    data=buffer,
                    file_name=f"{topic}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

if __name__ == "__main__":
    main()
